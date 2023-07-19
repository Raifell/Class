import docx
from docx.shared import Pt

num = input('Enter new integer: ')


def replace_string(filename, num):
    doc = docx.Document(filename)
    count = []
    number = ''
    drive = ''


    for p in doc.paragraphs:
        count.append(p.text)

    for x in count:
        if x.isdigit():
            number = x

    for x in range(len(count)):
        if number in count[x]:
            drive = count[x]

    serp = drive
    print(serp)
    ind = drive.find('=')
    new_drive = drive[:ind+1]
    dreben = str(num)+'*50%='
    print(new_drive)

    drive = int(num)*(int(num)/100*50)
    drive = str(drive)
    result = drive[:-2]
    print(result)

    aa = 0
    for p in doc.paragraphs:
        count.append(p.text)
        if number in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if number in inline[i].text:
                    text = inline[i].text.replace(number, num)
                    inline[i].text = text
                    aa += 1
                    print(p.text, 1)
            if aa > 0:
                break
    for p in doc.paragraphs:
        count.append(p.text)
        if number in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if serp in inline[i].text:
                    print(11, inline[i].text == serp, dreben+result)
                    text = inline[i].text.replace(serp, dreben+result)
                    inline[i].text = text

    # print(number)
    # print(drive)


    doc.save('dest1.docx')
    return 1


replace_string('example_new.docx', num)
