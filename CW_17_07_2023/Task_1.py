import docx
from docx.shared import Pt
import requests
from bs4 import BeautifulSoup

URL = 'https://loremipsum.io/'
page = requests.get(URL)

html = page.text
soup = BeautifulSoup(html, features='html.parser')

text = soup.find('div', {'class': 'page-section__generator g12-xs g11-md g10-lg g8-xl'})
result = text.find('blockquote', {'class': 'page-generator__lorem'})
result = result.text

count = result.replace(', ', ',|')
count = count.replace('. ', '.|')
count = count.split('|')
sep = ' '

doc = docx.Document()
doc.add_paragraph('Classwork', 'Title').alignment = 1
p1 = doc.add_paragraph('    ' + count[0] + sep + count[1] + sep)
p1.add_run(count[2] + sep).bold = True

p2 = doc.add_paragraph('    ' + count[3] + sep + count[4] + sep)
p2.add_run(count[5] + sep).bold = True
p2.add_run(count[6] + sep)
p2.add_run(count[7] + sep)

doc.save('example.docx')

doc = docx.Document('example.docx')

bolds = []
for para in doc.paragraphs:
    for run in para.runs:
        if run.bold:
            bolds.append(run.text)

doc = docx.Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(14)

doc.add_paragraph('Result', 'Title').alignment = 1
ps = doc.add_paragraph('')

xs = ps.add_run(' '.join(bolds))
xs.font.size = Pt(26)
xs.font.italic = True

doc.save('example_new.docx')
